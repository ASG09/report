from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from openai import OpenAI
from pathlib import Path
from tempfile import NamedTemporaryFile
import pdfplumber

app = FastAPI()

# Allow CORS for the frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # React frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

client = OpenAI()


@app.get("/")
async def root():
    return {"message": "Backend is running!"}


@app.post("/create-report")
async def create_report(cv: UploadFile = File(...), template: UploadFile = File(...)):

    model = "gpt-4o-mini"

    # Extract text from the CV file
    if cv.content_type == "application/pdf":
        try:
            with NamedTemporaryFile(delete=False, suffix=".pdf") as temp_cv:
                temp_cv.write(await cv.read())
                temp_cv_path = Path(temp_cv.name)

            cv_text = extract_text_from_pdf(temp_cv_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to process CV PDF: {e}")
    else:
        raise HTTPException(status_code=400, detail="CV must be in PDF format.")

    # Extract text from the template file
    if template.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            with NamedTemporaryFile(delete=False, suffix=".docx") as temp_template:
                temp_template.write(await template.read())
                temp_template_path = Path(temp_template.name)

            template_text = extract_text_from_docx(temp_template_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to process template DOCX: {e}")
    else:
        raise HTTPException(status_code=400, detail="Template must be in DOCX format.")

    # Use OpenAI API to generate the report based on extracted texts
    try:
        prompt = (
            "Using the information extracted from the following CV and template, generate a report with the template requested information:\n\n"
            "CV Content:\n" + cv_text + "\n\n"
            "Template Content:\n" + template_text
        )
        chat_completion = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
        )
        report_text = chat_completion.choices[0].message.content
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating report with OpenAI: {e}")

    # Save the generated report to a DOCX file
    try:
        with NamedTemporaryFile(delete=False, suffix=".docx") as temp_report:
            report_path = Path(temp_report.name)
            document = Document()
            document.add_heading("Generated Report", level=1)
            document.add_paragraph(report_text)
            document.save(report_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save report: {e}")

    # Return the generated DOCX file
    return FileResponse(
        report_path,
        filename="generated_report.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def extract_text_from_docx(docx_path: Path) -> str:
    """Extracts and returns all text from a DOCX file."""
    doc = Document(docx_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extracts and returns all text from a PDF file."""
    full_text = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            full_text.append(page.extract_text())
    return "\n".join(full_text)
